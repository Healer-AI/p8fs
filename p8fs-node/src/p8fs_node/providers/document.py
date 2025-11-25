"""Document content provider implementation."""

import logging
from pathlib import Path
from typing import Any, List, Optional

from p8fs_node.models.content import (
    ContentChunk,
    ContentMetadata,
    ContentProvider,
    ContentType,
)
from p8fs_node.providers.mixins import BaseProviderMixin, PlaceholderProviderMixin
from p8fs_node.utils.text import TextChunker, clean_text

logger = logging.getLogger(__name__)


class DocumentContentProvider(BaseProviderMixin, ContentProvider):
    """Content provider for document files (DOCX, ODT, RTF)."""

    @property
    def supported_types(self) -> list[ContentType]:
        """Return list of supported content types."""
        return [ContentType.DOCUMENT]

    @property
    def provider_name(self) -> str:
        """Return the name of the provider."""
        return "document_provider"

    async def extract_text(self, content_path: str | Path) -> str:
        """
        Extract raw text content from document file.
        
        Args:
            content_path: Path to the document file
            
        Returns:
            Raw text content extracted from the document
        """
        file_path = Path(content_path)

        if not file_path.exists():
            raise FileNotFoundError(f"Document file not found: {file_path}")

        try:
            if file_path.suffix.lower() == ".docx":
                text_content, _ = await self._extract_docx_content(file_path)
            elif file_path.suffix.lower() == ".odt":
                text_content, _ = await self._extract_odt_content(file_path)
            elif file_path.suffix.lower() == ".rtf":
                text_content, _ = await self._extract_rtf_content(file_path)
            else:
                return f"[Unsupported document format: {file_path.suffix}]"

            return text_content or ""
            
        except Exception as e:
            logger.error(f"Error extracting text from document {file_path}: {e}")
            return f"[Document text extraction failed: {str(e)}]"

    async def to_markdown_chunks(
        self, content_path: str | Path, extended: bool = False, **options: Any
    ) -> list[ContentChunk]:
        """Convert document content to semantic markdown chunks."""
        from p8fs_node.processors import MarkdownProcessor

        file_path = Path(content_path)

        if not file_path.exists():
            raise FileNotFoundError(f"Document file not found: {file_path}")

        # Extract text based on file type
        try:
            if file_path.suffix.lower() == ".docx":
                text_content, extraction_metadata = await self._extract_docx_content(
                    file_path, **options
                )
            elif file_path.suffix.lower() == ".odt":
                text_content, extraction_metadata = await self._extract_odt_content(
                    file_path, **options
                )
            elif file_path.suffix.lower() == ".rtf":
                text_content, extraction_metadata = await self._extract_rtf_content(
                    file_path, **options
                )
            else:
                # Fallback to placeholder for unsupported formats
                return self._create_placeholder_chunks(file_path, ContentType.DOCUMENT)

            if not text_content.strip():
                logger.warning(f"No text content extracted from {file_path}")
                return self._create_placeholder_chunks(file_path, ContentType.DOCUMENT)

            # Use centralized MarkdownProcessor
            chunks, _ = await MarkdownProcessor.process(
                text=text_content,
                source_file=file_path,
                content_type=ContentType.DOCUMENT,
                extraction_method=f"document_{file_path.suffix.lower().replace('.', '')}",
                provider_metadata=extraction_metadata,
                **options,
            )

            return chunks

        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return self._create_placeholder_chunks(file_path, ContentType.DOCUMENT)

    async def _extract_docx_content(
        self,
        file_path: Path,
        preserve_formatting: bool = False,
        extract_tables: bool = True,
        **options: Any,
    ) -> tuple[str, dict]:
        """Extract content from DOCX files."""
        try:
            from docx import Document

            doc = Document(file_path)
            content_parts = []
            has_tables = False
            has_images = False

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    if preserve_formatting:
                        # TODO: Implement formatting preservation
                        content_parts.append(paragraph.text)
                    else:
                        content_parts.append(paragraph.text)

            # Extract tables if requested
            if extract_tables and doc.tables:
                has_tables = True
                for table in doc.tables:
                    content_parts.append("\n## Table\n")
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            content_parts.append(f"| {row_text} |")
                    content_parts.append("")

            # Check for images (basic detection)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    has_images = True
                    break

            # Extract document properties
            props = doc.core_properties
            metadata = {
                "source": str(file_path),
                "file_type": "docx",
                "author": props.author or None,
                "title": props.title or file_path.stem,
                "created": props.created.isoformat() if props.created else None,
                "modified": props.modified.isoformat() if props.modified else None,
                "has_tables": has_tables,
                "has_images": has_images,
                "word_count": len(" ".join(content_parts).split()),
            }

            return "\n\n".join(content_parts), metadata

        except ImportError:
            logger.warning("python-docx not available, using docx2txt fallback")
            return await self._extract_docx_simple(file_path)
        except Exception as e:
            logger.error(f"Error extracting DOCX content: {e}")
            raise

    async def _extract_docx_simple(self, file_path: Path) -> tuple[str, dict]:
        """Simple DOCX extraction using docx2txt."""
        try:
            import docx2txt

            text = docx2txt.process(file_path)
            metadata = {
                "source": str(file_path),
                "file_type": "docx",
                "extraction_method": "docx2txt",
                "word_count": len(text.split()) if text else 0,
            }

            return text or "", metadata

        except ImportError:
            logger.error("Neither python-docx nor docx2txt available")
            raise ImportError("DOCX processing libraries not available")

    async def _extract_odt_content(
        self, file_path: Path, **options: Any
    ) -> tuple[str, dict]:
        """Extract content from ODT files."""
        try:
            from odf.opendocument import load
            from odf.text import P

            doc = load(file_path)
            content_parts = []

            # Extract paragraphs
            for paragraph in doc.getElementsByType(P):
                text = str(paragraph)
                if text.strip():
                    content_parts.append(text)

            metadata = {
                "source": str(file_path),
                "file_type": "odt",
                "extraction_method": "odfpy",
                "word_count": len(" ".join(content_parts).split()),
            }

            return "\n\n".join(content_parts), metadata

        except ImportError:
            logger.error("odfpy not available for ODT processing")
            raise ImportError("ODT processing library not available")
        except Exception as e:
            logger.error(f"Error extracting ODT content: {e}")
            raise

    async def _extract_rtf_content(
        self, file_path: Path, **options: Any
    ) -> tuple[str, dict]:
        """Extract content from RTF files."""
        try:
            from striprtf.striprtf import rtf_to_text

            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                rtf_content = f.read()

            text = rtf_to_text(rtf_content)

            metadata = {
                "source": str(file_path),
                "file_type": "rtf",
                "extraction_method": "striprtf",
                "word_count": len(text.split()) if text else 0,
            }

            return text or "", metadata

        except ImportError:
            logger.error("striprtf not available for RTF processing")
            raise ImportError("RTF processing library not available")
        except Exception as e:
            logger.error(f"Error extracting RTF content: {e}")
            raise

    async def to_metadata(
        self, content_path: str | Path, **options: Any
    ) -> ContentMetadata:
        """Extract metadata from document using MarkdownProcessor."""
        from p8fs_node.processors import MarkdownProcessor

        file_path = Path(content_path)

        if not file_path.exists():
            raise FileNotFoundError(f"Document file not found: {file_path}")

        try:
            # Extract text and get metadata from MarkdownProcessor
            if file_path.suffix.lower() == ".docx":
                text_content, extraction_metadata = await self._extract_docx_content(
                    file_path, **options
                )
                extraction_method = "document_docx"
            elif file_path.suffix.lower() == ".odt":
                text_content, extraction_metadata = await self._extract_odt_content(
                    file_path, **options
                )
                extraction_method = "document_odt"
            elif file_path.suffix.lower() == ".rtf":
                text_content, extraction_metadata = await self._extract_rtf_content(
                    file_path, **options
                )
                extraction_method = "document_rtf"
            else:
                # Basic metadata for unsupported formats
                from p8fs_node.utils.text import FileUtils

                stats = FileUtils.get_file_stats(file_path)

                return ContentMetadata(
                    title=stats["stem"],
                    file_path=str(file_path),
                    file_size=stats["size"],
                    created_date=stats["created"],
                    modified_date=stats["modified"],
                    content_type=ContentType.DOCUMENT,
                    extraction_method="unsupported_format",
                    properties={
                        "file_hash": stats["hash"],
                        "original_filename": stats["name"],
                        "file_extension": stats["suffix"],
                        "supported": False,
                    },
                )

            if not text_content.strip():
                # Return basic metadata for empty documents
                return ContentMetadata(
                    title=file_path.stem,
                    file_path=str(file_path),
                    file_size=file_path.stat().st_size,
                    content_type=ContentType.DOCUMENT,
                    extraction_method=f"{extraction_method}_empty",
                    word_count=0,
                    properties={"error": "No text content extracted"},
                )

            # Use MarkdownProcessor to get unified metadata
            _, metadata = await MarkdownProcessor.process(
                text=text_content,
                source_file=file_path,
                content_type=ContentType.DOCUMENT,
                extraction_method=extraction_method,
                provider_metadata=extraction_metadata,
                **options,
            )

            return metadata

        except Exception as e:
            logger.error(f"Error extracting document metadata: {e}")
            raise

    async def _extract_docx_metadata(self, file_path: Path) -> dict:
        """Extract metadata from DOCX file."""
        from p8fs_node.models.content import ContentMetadata
        from p8fs_node.utils.text import FileUtils

        try:
            from docx import Document

            doc = Document(file_path)
            props = doc.core_properties
            stats = FileUtils.get_file_stats(file_path)

            # Count paragraphs and tables
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            table_count = len(doc.tables)

            # Check for images
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1

            return ContentMetadata(
                title=props.title or stats["stem"],
                file_path=str(file_path),
                file_size=stats["size"],
                created_date=props.created or stats["created"],
                modified_date=props.modified or stats["modified"],
                content_type="document",
                extraction_method="python-docx",
                properties={
                    "file_hash": stats["hash"],
                    "original_filename": stats["name"],
                    "file_extension": stats["suffix"],
                    "author": props.author,
                    "subject": props.subject,
                    "paragraph_count": paragraph_count,
                    "table_count": table_count,
                    "image_count": image_count,
                    "has_tables": table_count > 0,
                    "has_images": image_count > 0,
                },
            )

        except ImportError:
            return await self._extract_basic_metadata(file_path, "docx")

    async def _extract_odt_metadata(self, file_path: Path) -> dict:
        """Extract metadata from ODT file."""
        try:
            from odf.opendocument import load
            from odf.text import P

            doc = load(file_path)

            # Count paragraphs
            paragraph_count = len(doc.getElementsByType(P))

            return await self._extract_basic_metadata(
                file_path,
                "odt",
                {
                    "paragraph_count": paragraph_count,
                    "extraction_method": "odfpy",
                },
            )

        except ImportError:
            return await self._extract_basic_metadata(file_path, "odt")

    async def _extract_rtf_metadata(self, file_path: Path) -> dict:
        """Extract metadata from RTF file."""
        return await self._extract_basic_metadata(file_path, "rtf")

    async def _extract_basic_metadata(
        self, file_path: Path, file_type: str, extra_props: dict = None
    ) -> dict:
        """Extract basic metadata for any document type."""
        from p8fs_node.models.content import ContentMetadata
        from p8fs_node.utils.text import FileUtils

        stats = FileUtils.get_file_stats(file_path)

        properties = {
            "file_hash": stats["hash"],
            "original_filename": stats["name"],
            "file_extension": stats["suffix"],
        }

        if extra_props:
            properties.update(extra_props)

        return ContentMetadata(
            title=stats["stem"],
            file_path=str(file_path),
            file_size=stats["size"],
            created_date=stats["created"],
            modified_date=stats["modified"],
            content_type="document",
            extraction_method=f"{file_type}_basic",
            properties=properties,
        )

    def _create_placeholder_chunks(
        self, file_path: Path, content_type: ContentType
    ) -> list[ContentChunk]:
        """Create placeholder chunks when processing fails."""
        placeholder_content = f"""# {file_path.name}

This is a placeholder for {content_type.value} content.

**File**: {file_path.name}
**Type**: {content_type.value}
**Provider**: {self.provider_name}

## Processing Status

The document could not be processed due to missing dependencies or errors.

*Note: This is placeholder content. Install required libraries (python-docx, docx2txt, odfpy, striprtf) for full document processing.*"""

        chunk = ContentChunk(
            id=f"{file_path.stem}_placeholder_chunk_0",
            content=placeholder_content,
            chunk_type="placeholder",
            position=0,
            metadata={
                "is_placeholder": True,
                "content_type": content_type.value,
                "provider": self.provider_name,
                "source": str(file_path),
                "file_type": file_path.suffix.lower().replace(".", ""),
            },
        )

        return [chunk]
